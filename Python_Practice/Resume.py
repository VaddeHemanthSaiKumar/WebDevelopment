from docx import Document

# Create a new Word document
doc = Document()

# Resume content
resume_text = """
VADDE HEMANTH SAI KUMAR
Hyderabad, Telangana, India | +91-9550060945 | saikumarjoey123@gmail.com | LinkedIn: linkedin.com/in/vadde-hemanth-sai-kumar

Professional Summary
Entry-level Associate Graduate with strong problem-solving and communication skills, focused on quality delivery and working closely with senior professionals. Experienced in SOP/runbook documentation, status reporting, data validation, and issue tracking. Eager to learn quickly, research solutions, and contribute to client-focused outcomes at NTT DATA.

Core Skills
- Programming & Scripting: Python; automation scripts for reporting and data processing
- Databases & Analysis: SQL, data validation, discrepancy checks, reporting
- Tools & Documentation: Git, Jupyter Notebook, Excel/Google Sheets, Confluence/Docs, presentation slides
- Delivery & Quality: Requirements understanding, defect identification, issue tracking, root-cause notes, status reporting
- Professional: Communication, stakeholder coordination, prioritization, time management, collaboration, continuous learning, adaptability

Education
B.Tech, Mechanical Engineering — Sree Vidyanikethan Engineering College | CGPA: 8.43/10 | 2024

Certifications
IBM Cybersecurity Analyst | Python Full Stack Development (JSpiders)

Projects & Technical Experience
Employee Data Management System (Python, SQL)
- Developed a Python system for employee records and automated report generation; improved processing efficiency by 30%
- Implemented SQL for data insert/update/retrieve and validation; documented corrections
- Authored SOPs and execution checklists; prepared weekly status reports and coordinated issue resolution

Process Documentation & Reporting Kit (Docs + Excel)
- Standardized runbook/templates for requirements, meeting notes, and status reporting
- Consolidated open questions, actions, and risks into a tracker for timely follow-ups

Internship Experience
Mechanical Intern — TTD Workshop, Tirupati
- Assisted in workflow optimization and record-keeping; maintained trackers and escalation notes
- Coordinated with internal teams and senior supervisors, producing clear documentation and checklists

Professional Attributes
Analytical mindset, fast learner, attentive to detail, organized planning and time management, team collaboration, committed to continuous improvement and quality delivery
"""

# Add each paragraph to the document
for paragraph in resume_text.strip().split("\n\n"):
    doc.add_paragraph(paragraph.strip())
    doc.add_paragraph("")  # Add spacing between sections

# Save the document
doc.save("NTT_DATA_Associate_Graduate_Resume_One_Page.docx")
print("Resume saved as NTT_DATA_Associate_Graduate_Resume_One_Page.docx")
